"""Document text extraction — PDF and DOCX parsing for tender intelligence.

Extracts raw text content from downloaded tender documents so the AI
analysis engine can work with the full scope and requirements.
"""

from __future__ import annotations

import os
from pathlib import Path

from src.core.logging import get_logger

logger = get_logger(__name__)


def extract_text(file_path: str) -> str:
    """Extract text content from a file based on its extension.

    Supported formats: PDF, DOCX, DOC (as plain text), TXT.
    Returns empty string on failure.
    """
    p = Path(file_path)
    if not p.exists():
        logger.warning("File not found: %s", file_path)
        return ""

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(p)
        elif ext in (".docx", ".doc"):
            return _extract_docx(p)
        elif ext in (".txt", ".csv"):
            return p.read_text(encoding="utf-8", errors="replace")
        else:
            logger.info("Unsupported file type %s for text extraction", ext)
            return ""
    except Exception as exc:
        logger.error("Text extraction failed for %s: %s", file_path, exc)
        return ""


def extract_text_from_directory(dir_path: str) -> dict[str, str]:
    """Extract text from all supported files in a directory.

    Returns dict mapping filename → extracted text.
    """
    results: dict[str, str] = {}
    d = Path(dir_path)
    if not d.is_dir():
        return results

    for f in sorted(d.iterdir()):
        if f.is_file() and f.suffix.lower() in (".pdf", ".docx", ".doc", ".txt"):
            text = extract_text(str(f))
            if text.strip():
                results[f.name] = text
                logger.info("Extracted %d chars from %s", len(text), f.name)

    return results


def count_pdf_pages(file_path: str) -> int | None:
    """Return the page count for a PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        return len(reader.pages)
    except Exception:
        return None


def _extract_pdf(path: Path) -> str:
    """Extract text from all pages of a PDF."""
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(text)
    full = "\n\n".join(pages)
    logger.info("PDF %s: %d pages, %d chars extracted", path.name, len(reader.pages), len(full))
    return full


def _extract_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    from docx import Document

    doc = Document(str(path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    full = "\n".join(paragraphs)
    logger.info("DOCX %s: %d paragraphs, %d chars extracted", path.name, len(paragraphs), len(full))
    return full
